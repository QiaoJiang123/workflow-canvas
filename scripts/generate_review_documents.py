from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "public" / "review-documents"
NODE_OUTPUT_DIR = OUTPUT_DIR / "approval-chain"


DOCUMENTS = {
    "underwriting-change-request.pdf": {
        "title": "Underwriting Change Request",
        "owner": "Riley Morgan",
        "due": "2026-08-05",
        "sections": [
            ("Summary", "Request approval to route medium-risk commercial property submissions through an AI-assisted triage flow."),
            ("Requested Change", "Add OCR, extraction, triage rules, human approval, and governance evidence storage."),
            ("Controls", "No automated binding decision. All recommendations include evidence, reviewer ownership, overrides, and rollback notes."),
        ],
    },
    "compliance-checklist.pdf": {
        "title": "Compliance Checklist",
        "owner": "Maya Chen",
        "due": "2026-08-05",
        "sections": [
            ("Required Evidence", "Business owner, impacted segments, fairness risks, data sensitivity, retention, monitoring, and escalation thresholds."),
            ("Review Notes", "Confirm the change supports triage only and routes missing evidence, high uncertainty, and restricted classes to manual review."),
            ("Approval Conditions", "Approval status must stay visible; rejected items route to governance review; rollback plan is required."),
        ],
    },
    "legal-terms-review.pdf": {
        "title": "Legal Terms Review",
        "owner": "Jordan Patel",
        "due": "2026-08-07",
        "sections": [
            ("Scope", "Review customer-facing language, vendor obligations, retention commitments, third-party services, and audit evidence."),
            ("Questions", "Confirm no binding underwriting decision is automated and extracted facts are retained for approved purposes only."),
            ("Draft Finding", "Acceptable if release notes state the workflow supports internal triage and not binding coverage or pricing decisions."),
        ],
    },
}

NODE_DOCUMENTS = [
    {
        "slug": "change-request-packet",
        "title": "Change Request Packet SOP",
        "owner": "Underwriting Operations",
        "purpose": "Confirm the request packet includes the decision context, impacted systems, and evidence required before triage.",
        "checks": ["Business owner and release owner named", "Review documents attached", "Rollback and monitoring notes present"],
    },
    {
        "slug": "completeness-check",
        "title": "Completeness Check SOP",
        "owner": "AI Governance",
        "purpose": "Validate that required approval evidence is present before assigning reviewers.",
        "checks": ["Risk tier recorded", "Affected policies listed", "Evidence packet and due dates complete"],
    },
    {
        "slug": "assign-compliance",
        "title": "Compliance Assignment SOP",
        "owner": "Compliance",
        "purpose": "Assign a compliance reviewer and define the evidence they must inspect.",
        "checks": ["Compliance reviewer selected", "Due date confirmed", "Checklist link attached"],
    },
    {
        "slug": "compliance-review",
        "title": "Compliance Review Packet",
        "owner": "Compliance",
        "purpose": "Review fairness, auditability, retention, disclosure, and monitoring obligations.",
        "checks": ["Fairness notes reviewed", "Audit evidence sufficient", "Open issues captured"],
    },
    {
        "slug": "compliance-approval",
        "title": "Compliance Approval SOP",
        "owner": "Compliance",
        "purpose": "Record formal compliance approval or route unresolved issues to escalation.",
        "checks": ["Checklist complete", "No high-risk open item", "Approval decision recorded"],
    },
    {
        "slug": "assign-legal",
        "title": "Legal Assignment SOP",
        "owner": "Legal",
        "purpose": "Assign legal review and clarify contract, notice, and retention questions.",
        "checks": ["Legal reviewer selected", "Terms review attached", "Due date confirmed"],
    },
    {
        "slug": "legal-review",
        "title": "Legal Review Packet",
        "owner": "Legal",
        "purpose": "Review customer-facing language, vendor obligations, retention, and non-binding decision controls.",
        "checks": ["Customer wording reviewed", "Vendor terms checked", "Retention language acceptable"],
    },
    {
        "slug": "legal-approval",
        "title": "Legal Approval SOP",
        "owner": "Legal",
        "purpose": "Record legal approval conditions before the business owner makes the final decision.",
        "checks": ["Conditions documented", "No blocking legal issue", "Approval decision recorded"],
    },
    {
        "slug": "business-owner-approval",
        "title": "Business Owner Approval SOP",
        "owner": "Underwriting Operations",
        "purpose": "Confirm the business owner accepts the release window, residual risk, and operating process.",
        "checks": ["Prior approvals complete", "Release window accepted", "Residual risk accepted"],
    },
    {
        "slug": "escalate-exception",
        "title": "Exception Escalation SOP",
        "owner": "AI Governance",
        "purpose": "Route rejected, overdue, or disputed approval items to governance review.",
        "checks": ["Exception reason logged", "Reviewer notified", "Next decision meeting assigned"],
    },
    {
        "slug": "notify-request-owner",
        "title": "Approval Notification SOP",
        "owner": "Release Coordination",
        "purpose": "Notify request owners and approvers of the final decision and next steps.",
        "checks": ["Decision summary sent", "Release owner notified", "Evidence links included"],
    },
    {
        "slug": "approval-audit-record",
        "title": "Approval Audit Record SOP",
        "owner": "AI Governance",
        "purpose": "Store the approval packet, decisions, evidence links, and reviewer trail for audit.",
        "checks": ["All decisions archived", "Documents linked", "Audit owner confirmed"],
    },
]


def build_pdf(filename: str, document: dict[str, object]) -> None:
    styles = getSampleStyleSheet()
    path = OUTPUT_DIR / filename
    story = [Paragraph(str(document["title"]), styles["Title"]), Spacer(1, 0.18 * inch)]
    meta = Table(
        [["Owner", document["owner"]], ["Review due", document["due"]], ["Document type", "Dummy review packet"]],
        colWidths=[1.2 * inch, 5.6 * inch],
    )
    meta.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f1f5f9")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#334155")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.extend([meta, Spacer(1, 0.24 * inch)])
    for heading, body in document["sections"]:
        story.append(Paragraph(str(heading), styles["Heading2"]))
        story.append(Paragraph(str(body), styles["BodyText"]))
        story.append(Spacer(1, 0.16 * inch))
    story.append(Paragraph("Generated dummy document for Flow Canvas approval-chain testing.", styles["Italic"]))
    SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.7 * inch, leftMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch).build(story)


def build_node_pdf(item: dict[str, object]) -> None:
    styles = getSampleStyleSheet()
    path = NODE_OUTPUT_DIR / f"{item['slug']}.pdf"
    story = [Paragraph(str(item["title"]), styles["Title"]), Spacer(1, 0.14 * inch)]
    story.append(Paragraph(f"Owner: {item['owner']}", styles["BodyText"]))
    story.append(Paragraph("Document type: SOP / review packet", styles["BodyText"]))
    story.append(Spacer(1, 0.16 * inch))
    story.append(Paragraph("Purpose", styles["Heading2"]))
    story.append(Paragraph(str(item["purpose"]), styles["BodyText"]))
    story.append(Spacer(1, 0.12 * inch))
    rows = [["Review check", "Expected evidence"]]
    rows.extend([[check, "Reviewer initials or note"] for check in item["checks"]])
    table = Table(rows, colWidths=[3.5 * inch, 2.8 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0b2545")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.extend([table, Spacer(1, 0.16 * inch)])
    story.append(Paragraph("Generated dummy document for Flow Canvas approval-chain node review.", styles["Italic"]))
    SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.7 * inch, leftMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.7 * inch).build(story)


def build_node_docx(item: dict[str, object]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if style_name == "Heading 3" else 14)
        style.paragraph_format.space_after = Pt(7 if style_name == "Heading 2" else 5)

    title = document.add_paragraph()
    title_format = title.paragraph_format
    title_format.space_after = Pt(8)
    title_run = title.add_run(str(item["title"]))
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor.from_string("0B2545")

    metadata = document.add_paragraph()
    metadata.add_run("Owner: ").bold = True
    metadata.add_run(str(item["owner"]))
    metadata.add_run("    Type: ").bold = True
    metadata.add_run("SOP / review packet")

    document.add_heading("Purpose", level=1)
    document.add_paragraph(str(item["purpose"]))

    document.add_heading("Review Checklist", level=1)
    for check in item["checks"]:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(str(check))

    document.add_heading("Approver Notes", level=1)
    for label in ["Decision", "Conditions", "Follow-up owner"]:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run("____________________________________________________________")

    footer = section.footer.paragraphs[0]
    footer.text = "Flow Canvas dummy SOP"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor.from_string("64748B")

    document.save(NODE_OUTPUT_DIR / f"{item['slug']}.docx")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    NODE_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for filename, document in DOCUMENTS.items():
        build_pdf(filename, document)
    for item in NODE_DOCUMENTS:
        build_node_pdf(item)
        build_node_docx(item)


if __name__ == "__main__":
    main()
